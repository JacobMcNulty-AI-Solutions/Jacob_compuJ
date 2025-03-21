import io
import logging
from typing import Tuple, Optional, Dict, Any
import PyPDF2
import docx
import os
from app.core.exceptions import InvalidFileFormat

# Set up logging
logger = logging.getLogger(__name__)

class DocumentProcessingError(Exception):
    """Base class for document processing errors"""
    def __init__(self, message: str, details: Dict[str, Any] = None):
        self.message = message
        self.details = details or {}
        super().__init__(self.message)

class PDFExtractionError(DocumentProcessingError):
    """Raised when PDF text extraction fails"""
    pass

class DocxExtractionError(DocumentProcessingError):
    """Raised when DOCX text extraction fails"""
    pass

class UnsupportedFileTypeError(DocumentProcessingError):
    """Raised when file type is not supported"""
    pass

class CorruptedFileError(DocumentProcessingError):
    """Raised when file is corrupted and cannot be processed"""
    pass

class DocumentProcessor:
    """Service for processing document files (text extraction, etc.)"""
    
    # Maximum file size in bytes (10MB)
    MAX_FILE_SIZE = 10 * 1024 * 1024
    
    # Supported content types
    SUPPORTED_CONTENT_TYPES = {
        "text/plain": [".txt"],
        "application/pdf": [".pdf"],
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": [".docx"],
        "application/msword": [".doc"]
    }
    
    @staticmethod
    def diagnose_pdf(file_content: bytes, filename: str = "") -> dict:
        """
        Debug function to diagnose issues with PDF files
        
        Purpose:
        - Identify common problems with PDF files
        - Provide detailed diagnostic information
        - Help troubleshoot PDF processing failures
        
        Args:
            file_content: Raw bytes of the PDF file
            filename: Original filename 
            
        Returns:
            Dictionary with diagnostic information
        """
        diagnostic = {
            "filename": filename,
            "file_size": len(file_content),
            "checks": {}
        }
        
        # Check 1: PDF header
        try:
            has_pdf_header = file_content.startswith(b'%PDF-')
            diagnostic["checks"]["has_pdf_header"] = has_pdf_header
            if has_pdf_header:
                pdf_version = file_content[5:8].decode('ascii', errors='ignore')
                diagnostic["checks"]["pdf_version"] = pdf_version
        except Exception as e:
            diagnostic["checks"]["header_error"] = str(e)
        
        # Check 2: Try reading with PyPDF2
        try:
            with io.BytesIO(file_content) as pdf_file:
                try:
                    reader = PyPDF2.PdfReader(pdf_file)
                    diagnostic["checks"]["pypdf2_read"] = "success"
                    diagnostic["checks"]["is_encrypted"] = reader.is_encrypted
                    diagnostic["checks"]["page_count"] = len(reader.pages)
                    
                    # Try to read first page
                    if len(reader.pages) > 0:
                        try:
                            first_page = reader.pages[0]
                            text = first_page.extract_text()
                            diagnostic["checks"]["first_page_has_text"] = bool(text.strip())
                            diagnostic["checks"]["first_page_text_length"] = len(text)
                        except Exception as e:
                            diagnostic["checks"]["first_page_error"] = str(e)
                except PyPDF2.errors.PdfReadError as e:
                    diagnostic["checks"]["pypdf2_read"] = f"failed: {str(e)}"
                except Exception as e:
                    diagnostic["checks"]["pypdf2_read"] = f"unknown error: {str(e)}"
        except Exception as e:
            diagnostic["checks"]["io_error"] = str(e)
            
        return diagnostic
    
    @staticmethod
    def validate_file(file_content: bytes, content_type: str, filename: str) -> None:
        """
        Validate file size and type before processing
        
        Purpose:
        - Verify the file meets size constraints
        - Ensure file type is supported
        - Check that file extension matches content type
        - Prevent processing of invalid files early
        
        Args:
            file_content: Raw bytes of the file
            content_type: MIME type of the file
            filename: Name of the file
            
        Raises:
            DocumentProcessingError: If file size exceeds maximum allowed size
            UnsupportedFileTypeError: If content type is not supported
            UnsupportedFileTypeError: If file extension doesn't match content type
        """
        # Check file size
        if len(file_content) > DocumentProcessor.MAX_FILE_SIZE:
            raise DocumentProcessingError(
                f"File size exceeds maximum allowed size of {DocumentProcessor.MAX_FILE_SIZE / 1024 / 1024}MB",
                {"max_size_bytes": DocumentProcessor.MAX_FILE_SIZE, "actual_size_bytes": len(file_content)}
            )
        
        # Check content type
        if content_type not in DocumentProcessor.SUPPORTED_CONTENT_TYPES:
            raise UnsupportedFileTypeError(
                f"Unsupported content type: {content_type}",
                {"supported_types": list(DocumentProcessor.SUPPORTED_CONTENT_TYPES.keys())}
            )
        
        # Double-check file extension matches content type
        _, file_ext = os.path.splitext(filename.lower())
        if file_ext not in DocumentProcessor.SUPPORTED_CONTENT_TYPES.get(content_type, []):
            raise UnsupportedFileTypeError(
                f"File extension {file_ext} does not match content type {content_type}",
                {
                    "extension": file_ext,
                    "content_type": content_type,
                    "expected_extensions": DocumentProcessor.SUPPORTED_CONTENT_TYPES.get(content_type, [])
                }
            )
    
    @staticmethod
    def extract_text(file_content: bytes, content_type: str, filename: str = "") -> Tuple[str, Optional[str]]:
        """
        Extract text from document based on content type
        
        Purpose:
        - Extract readable text content from various document formats
        - Handle different file types appropriately (PDF, DOCX, TXT)
        - Provide error information if extraction fails
        
        Args:
            file_content: Raw bytes of the file
            content_type: MIME type of the file
            filename: Original filename (for extension validation)
            
        Returns:
            Tuple containing: (extracted_text, error_message)
            where error_message is None if successful, or contains an error description if failed
        """
        try:
            # Validate file before processing
            try:
                DocumentProcessor.validate_file(file_content, content_type, filename)
            except DocumentProcessingError as e:
                logger.warning(f"File validation failed: {str(e)}", extra={"details": e.details})
                return "", str(e)
                
            # Plain text files
            if content_type == "text/plain":
                try:
                    return file_content.decode("utf-8"), None
                except UnicodeDecodeError:
                    error_msg = "Unable to decode text file. The file may be binary or use an unsupported encoding."
                    logger.warning(error_msg)
                    return "", error_msg
                
            # PDF files
            elif content_type == "application/pdf":
                try:
                    logger.info(f"Starting PDF processing for file: {filename}")
                    with io.BytesIO(file_content) as pdf_file:
                        # Check if file is encrypted/password protected
                        try:
                            logger.info("Initializing PDF reader")
                            reader = PyPDF2.PdfReader(pdf_file)
                            logger.info(f"PDF reader initialized, is_encrypted: {reader.is_encrypted}")
                            
                            if reader.is_encrypted:
                                error_msg = "PDF is encrypted and cannot be processed without a password."
                                logger.warning(error_msg)
                                return "", error_msg
                            
                            # Check if PDF has pages
                            if len(reader.pages) == 0:
                                logger.warning("PDF has no pages")
                                return "", "PDF has no pages"
                                
                            # Extract text from all pages
                            text = ""
                            logger.info(f"PDF has {len(reader.pages)} pages")
                            for i, page in enumerate(reader.pages):
                                try:
                                    logger.info(f"Extracting text from page {i+1}/{len(reader.pages)}")
                                    page_text = page.extract_text()
                                    if page_text:
                                        text += page_text + "\n"
                                        logger.debug(f"Extracted {len(page_text)} characters from page {i+1}")
                                    else:
                                        logger.warning(f"No text extracted from page {i+1}")
                                except Exception as page_err:
                                    logger.warning(f"Error extracting text from page {i}: {str(page_err)}", exc_info=True)
                                    # Continue with other pages
                            
                            # Check if we got any text
                            if not text.strip():
                                logger.warning("PDF appears to contain no extractable text (may be scanned document)")
                                return "", "PDF appears to contain no extractable text. The document may be scanned or contain only images."
                                
                            logger.info(f"Successfully extracted {len(text)} characters from PDF")
                            return text, None
                            
                        except PyPDF2.errors.PdfReadError as pdf_err:
                            error_msg = f"Invalid or corrupted PDF file: {str(pdf_err)}"
                            logger.warning(error_msg, exc_info=True)
                            
                            # Try to reset the file position and read again with a more lenient approach
                            try:
                                pdf_file.seek(0)
                                logger.info("Trying alternative PDF reading approach")
                                
                                # Check if it's even a PDF by looking at magic bytes
                                pdf_file.seek(0)
                                header = pdf_file.read(5)
                                if header != b'%PDF-':
                                    logger.error("File does not have PDF header bytes")
                                    return "", "Not a valid PDF file. Missing PDF header."
                                
                                # Try a more lenient approach - if the file has a PDF header but can't be read
                                # by PyPDF2, we'll attempt to extract at least some text by scanning for text strings
                                logger.info("PDF header found, trying basic text extraction")
                                pdf_file.seek(0)
                                pdf_content = pdf_file.read()
                                
                                # Very simple text extraction - just look for readable ASCII text
                                import re
                                text_chunks = re.findall(b'[a-zA-Z0-9 .,;:!?\'"\-+=/\\\\()\[\]{}]{4,}', pdf_content)
                                if text_chunks:
                                    extracted_text = b'\n'.join(text_chunks).decode('utf-8', errors='ignore')
                                    logger.info(f"Found {len(text_chunks)} text chunks using fallback method")
                                    if len(extracted_text) > 100:  # If we found a reasonable amount of text
                                        return extracted_text, None
                                        
                                # If we couldn't extract text or found too little, return the original error
                                pdf_file.seek(0)
                                return "", "The PDF file appears to be corrupted or in an unsupported format."
                            except Exception as e:
                                logger.error(f"Alternative PDF reading also failed: {str(e)}")
                                return "", "The PDF file could not be processed."
                            
                except Exception as pdf_err:
                    error_msg = f"PDF extraction error: {str(pdf_err)}"
                    logger.error(error_msg, exc_info=True)
                    return "", error_msg
                
            # Word documents
            elif "application/vnd.openxmlformats-officedocument.wordprocessingml" in content_type:
                try:
                    with io.BytesIO(file_content) as docx_file:
                        try:
                            doc = docx.Document(docx_file)
                            text = ""
                            # Extract text from paragraphs
                            for para in doc.paragraphs:
                                text += para.text + "\n"
                            
                            # Extract text from tables
                            for table in doc.tables:
                                for row in table.rows:
                                    for cell in row.cells:
                                        text += cell.text + " "
                                    text += "\n"
                            
                            # Check if we got any text
                            if not text.strip():
                                logger.warning("DOCX appears to contain no text")
                                return "", "DOCX document appears to contain no text."
                                
                            return text, None
                            
                        except Exception as docx_format_err:
                            error_msg = f"Invalid or corrupted DOCX file: {str(docx_format_err)}"
                            logger.warning(error_msg)
                            return "", error_msg
                            
                except Exception as docx_err:
                    error_msg = f"DOCX extraction error: {str(docx_err)}"
                    logger.error(error_msg, exc_info=True)
                    return "", error_msg
                
            else:
                error_msg = f"Unsupported content type: {content_type}"
                logger.warning(error_msg)
                return "", error_msg
                
        except Exception as e:
            error_msg = f"Error extracting text: {str(e)}"
            logger.error(error_msg, exc_info=True)
            return "", error_msg 